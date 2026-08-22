from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parent
SHOTS = ROOT / "screenshots"
FIELD = Path(r"C:\Users\selorm\Downloads")
OUT = ROOT / "BuildMate_Ghana_Regent_Capstone_Report.docx"

GREEN = RGBColor(20, 92, 57)
DARK = RGBColor(18, 37, 31)
MUTED = RGBColor(91, 105, 100)
LIGHT = "EAF4EF"
GOLD = RGBColor(174, 116, 24)

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.68)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal.font.size = Pt(10.2)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.16

for name, size, before, after in [
    ("Heading 1", 16, 14, 8),
    ("Heading 2", 13, 10, 5),
    ("Heading 3", 11.5, 7, 3),
]:
    style = styles[name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = GREEN if name != "Heading 3" else DARK
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

caption = styles["Caption"]
caption.font.name = "Times New Roman"
caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
caption.font.size = Pt(8.5)
caption.font.italic = True
caption.font.color.rgb = DARK
caption.paragraph_format.space_before = Pt(3)
caption.paragraph_format.space_after = Pt(5)
caption.paragraph_format.keep_with_next = True

if "Figure Explanation" not in styles:
    ex = styles.add_style("Figure Explanation", WD_STYLE_TYPE.PARAGRAPH)
else:
    ex = styles["Figure Explanation"]
ex.font.name = "Times New Roman"
ex._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
ex._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
ex.font.size = Pt(9.3)
ex.font.color.rgb = MUTED
ex.paragraph_format.space_after = Pt(5)
ex.paragraph_format.line_spacing = 1.1


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    tbl_pr.append(ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    r = p.add_run("BuildMate Ghana Capstone Report")
    r.font.name = "Times New Roman"
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED


add_footer(section)


def new_page():
    doc.add_page_break()


def title(text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(23)
    r.font.bold = True
    r.font.color.rgb = GREEN
    if subtitle:
        q = doc.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.paragraph_format.space_after = Pt(12)
        s = q.add_run(subtitle)
        s.font.name = "Times New Roman"
        s.font.size = Pt(12)
        s.font.italic = True
        s.font.color.rgb = MUTED


def h1(text):
    p = doc.add_paragraph(text, style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.right_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)


def h2(text):
    p = doc.add_paragraph(text, style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.right_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)


def para(text, bold_lead=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        a, b = text[: len(bold_lead)], text[len(bold_lead) :]
        p.add_run(a).bold = True
        p.add_run(b)
    else:
        p.add_run(text)
    return p


def compact_lines(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(line)


def callout(label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.add_run(label + ": ").bold = True
    p.add_run(text)


def add_picture(run, path, width, alt_text):
    shape = run.add_picture(str(path), width=width)
    properties = shape._inline.docPr
    properties.set("descr", alt_text)
    properties.set("title", alt_text)
    return shape


def figure(filename, number, caption_text, explanation, width=6.75):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    add_picture(p.add_run(), SHOTS / filename, Inches(width), caption_text)
    c = doc.add_paragraph(f"Figure 4.{number}: {caption_text}", style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    e = doc.add_paragraph(explanation, style="Figure Explanation")
    return e


def field_photo_grid(items):
    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for start in range(0, len(items), 3):
        image_cells = table.add_row().cells
        caption_cells = table.add_row().cells
        for offset in range(3):
            index = start + offset
            set_cell_margins(image_cells[offset], 35, 45, 20, 45)
            set_cell_margins(caption_cells[offset], 0, 45, 50, 45)
            if index >= len(items):
                continue
            filename, label, alt = items[index]
            image_p = image_cells[offset].paragraphs[0]
            image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_picture(image_p.add_run(), FIELD / filename, Inches(1.95), alt)
            caption_p = caption_cells[offset].paragraphs[0]
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_p.style = "Caption"
            caption_p.add_run(label)
    set_table_geometry(table, [3120, 3120, 3120])


def system_snapshot(filename, caption_text, explanation, width=5.55):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    add_picture(p.add_run(), SHOTS / filename, Inches(width), caption_text)
    c = doc.add_paragraph(caption_text, style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    e = doc.add_paragraph(explanation, style="Figure Explanation")
    e.alignment = WD_ALIGN_PARAGRAPH.CENTER


# Page 1: cover
logo = doc.add_paragraph()
logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo.paragraph_format.space_after = Pt(4)
add_picture(logo.add_run(), ROOT / "regentlogo1.png", Inches(1.35), "Regent University College of Science and Technology logo")
title("BUILDMATE GHANA", "Digitising Nana Attakorah II Ventures through a Secure Multi-Tenant Marketplace")
doc.add_paragraph().paragraph_format.space_after = Pt(6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FINAL CAPSTONE PROJECT REPORT")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = GOLD
doc.add_paragraph().paragraph_format.space_after = Pt(8)
for line in [
    "Programme: COMPUTER SCIENCE & ISS",
    "Course: E-COMMERCE",
    "Lecturer: Emmanuel Dzorkah (E-Commerce Lecturer)",
    "Institution: Regent University College of Science and Technology",
    "Academic Year: 2025/2026",
]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
group_label = doc.add_paragraph()
group_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
group_label.paragraph_format.space_before = Pt(4)
group_label.paragraph_format.space_after = Pt(2)
group_label.add_run("Project Group").bold = True
group = doc.add_table(rows=1, cols=2)
group.style = "Table Grid"
group.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, value in enumerate(["Student", "Index Number"]):
    group.rows[0].cells[i].text = value
    set_cell_shading(group.rows[0].cells[i], LIGHT)
    group.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    group.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for name, index_number in [
    ("David Selorm Gabion", "15240020"),
    ("Sualiha Nasomah Zakari", "11240016"),
    ("Charles Nanor", "15240047"),
    ("Shadrack Ofori", "15240042"),
    ("Mary Sarpong", "15240061"),
]:
    cells = group.add_row().cells
    cells[0].text = name
    cells[1].text = index_number
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
set_table_geometry(group, [6200, 3160])
for row in group.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
para("A capstone report submitted in partial fulfilment of the requirements for the award of the stated programme.", align=WD_ALIGN_PARAGRAPH.CENTER)

# Page 2
new_page()
h1("Declaration and Abstract")
h2("Declaration")
para("We, David Selorm Gabion, Sualiha Nasomah Zakari, Charles Nanor, Shadrack Ofori and Mary Sarpong, declare that this report describes the BuildMate Ghana system as inspected in the project repository, field study and running deployment. The work is submitted for the E-Commerce course in the Computer Science & ISS programme at Regent University College of Science and Technology.")
h2("Abstract")
para("This capstone project uses Nana Attakorah II Ventures as the selected local building-materials business and places it on BuildMate Ghana, a multi-tenant construction procurement platform. The project responds to the business's limited digital presence and to wider problems of fragmented sourcing, uncertain stock availability, informal quotations and weak fulfilment visibility. BuildMate gives Nana Attakorah II Ventures its own protected supplier-organisation tenant for products, branches, inventory, staff, quotations, orders and settlements while sharing a public marketplace with other independently isolated suppliers.")
para("The system uses Next.js 16 and React 19 for its application layer and Supabase/PostgreSQL for authentication, relational storage, row-level security and transactional workflows. Organisation membership, permissions and branch assignments isolate each supplier tenant so Nana Attakorah II Ventures can manage its operations without seeing or altering another supplier's private records. The completed catalogue now contains 15 active top-level categories, 36 active subcategories and 129 active canonical products. Verification recorded 312 passing automated tests, a clean TypeScript check, zero-warning ESLint execution and a successful production build. The study concludes that BuildMate provides a credible path for bringing the selected business online while preserving scalable marketplace governance. Field photographs, public application screens, an authenticated administration screen and repository verification provide the supporting evidence used in this report.")
h2("Keywords")
para("Construction procurement; e-commerce; marketplace; inventory ledger; role-based access control; Supabase; Next.js; Ghana.")

# Page 3
new_page()
h1("Table of Contents")
compact_lines([
    "Declaration and Abstract ........................................................ 2",
    "Chapter One: Introduction ...................................................... 4",
    "Chapter Two: Business Analysis ................................................. 6",
    "Chapter Three: System Design .................................................... 8",
    "Chapter Four: System Implementation ............................................ 10",
    "Chapter Five: Conclusion and Recommendations ................................... 17",
    "References and Project-Brief Alignment ....................................... 18",
])
h2("List of Figures")
compact_lines([
    "Figure 3.1: BuildMate Ghana Use Case Diagram .................................... 8",
    "Figures 4.1-4.14: Business Field-Study Photographs ......................... 10-12",
    "Figure 4.15: BuildMate Ghana Homepage .......................................... 13",
    "Figure 4.16: Building Materials Marketplace ................................... 13",
    "Figure 4.17: Construction Services Marketplace ................................ 14",
    "Figure 4.18: Material Calculators .............................................. 14",
    "Figure 4.19: Competitive Quotation Request Interface .......................... 15",
    "Figure 4.20: Shopping Cart and Fulfilment Interface ........................... 15",
    "Figure 4.21: Responsive Mobile Homepage ........................................ 16",
    "Figure 4.22: Platform Administration Dashboard ................................ 16",
])
h2("Abbreviations")
para("API - Application Programming Interface; COD - Cash on Delivery; GRN - Goods Received Note; RLS - Row Level Security; RPC - Remote Procedure Call; RBAC - Role-Based Access Control; SKU - Stock Keeping Unit; UI - User Interface.")

# Page 4
new_page()
h1("Chapter One: Introduction")
h2("1.1 Background")
para("Nana Attakorah II Ventures is the local building-materials business selected for this capstone project. The field study documented physical stock, face-to-face sales discussions, roadside merchandising and local transport. Without a structured online channel, customers depend heavily on visits, telephone calls and informal messages to discover products, confirm availability and arrange collection or delivery.")
para("BuildMate Ghana addresses this limitation by onboarding Nana Attakorah II Ventures as an independent supplier tenant on a shared construction marketplace. Public users can discover its eligible offers alongside other approved suppliers, while the business receives a protected organisation workspace for products, branch-level inventory, quotations, orders, staff and settlements. Customer, driver, service-provider and platform-administration workflows remain shared platform services, but private supplier records are isolated by tenant permissions.")
h2("1.2 Problem Statement")
para("Nana Attakorah II Ventures has a real physical presence and visible stock, but a limited functional online marketplace. Its existing sales environment does not consistently link public product discovery, current stock, formal quotations, order acknowledgement, delivery milestones and payment evidence. The solution must therefore bring the business online while keeping its organisation, branches, staff, prices, inventory and orders isolated from other suppliers using the same platform.")
h2("1.3 Aim")
para("The project aims to digitise Nana Attakorah II Ventures by onboarding it onto a secure, responsive and auditable multi-tenant platform through which customers can discover materials and the business can manage resulting commercial workflows.")
h2("1.4 Objectives")
compact_lines([
    "1. Provide a public catalogue with search, category filtering and supplier offers.",
    "2. Support quotation, cart, order, fulfilment and customer-confirmation workflows.",
    "3. Provide Nana Attakorah II Ventures with controlled tools for products, inventory, orders, staff and finance.",
    "4. Isolate Nana Attakorah II Ventures from other supplier tenants through organisation, permission and branch controls.",
    "5. Provide platform operations, reporting, support and auditable administrative review.",
])

# Page 5
new_page()
h1("Chapter One: Scope and Significance")
h2("1.5 Scope")
para("The implemented scope includes public marketplace browsing, service-provider discovery, material calculators, assisted quotation requests, shopping-cart fulfilment selection, customer and organisation workspaces, supplier management, branch-aware inventory, orders, delivery operations, settlement views, unified support tickets, provider verification and platform administration. The system supports cash on delivery and cash on pickup; it does not claim that every external payment, logistics or architectural-review integration is fully automated.")
h2("1.6 Significance")
para("For customers, the platform makes Nana Attakorah II Ventures discoverable and introduces structured transaction records. For the selected business, it separates selling price from unit cost, records immutable stock movements and supports organisation staff permissions. For BuildMate operations, it provides supplier approval, support, audit and reporting tools. The design demonstrates how one local business can gain an online presence inside a larger marketplace without losing ownership or privacy of its operational data.")
h2("1.7 Research Questions")
compact_lines([
    "1. How can a construction marketplace combine public discovery with protected multi-role operations?",
    "2. How can database controls preserve legitimate inventory availability and tenant isolation?",
    "3. How can responsive interfaces simplify procurement without hiding operational complexity?",
    "4. What evidence demonstrates that the implementation is buildable, testable and aligned with its objectives?",
])
h2("1.8 Limitations")
para("This report evaluates the current codebase and the live deployment available through 22 August 2026. The public catalogue exposes all 15 active top-level categories, while the product-offer marketplace may legitimately show no purchasable materials until approved suppliers publish branch-assigned, positively priced and eligible stock. The service directory returned no verified matching providers during the original evidence capture. The administrator dashboard was accessible at the start of capture, but the protected session expired before customer and supplier portal screenshots could be collected. No credentials, stock levels, warehouse assignments or publication states were requested, extracted or fabricated.")
callout("Assessment position", "Screenshots are treated as evidence of the observed system state, while repository inspection and automated tests provide evidence for protected workflows that could not be re-opened visually.")

# Page 6
new_page()
h1("Chapter Two: Business Analysis")
h2("2.1 Description of the Selected Business")
para("Nana Attakorah II Ventures is the selected local building-materials business. Field observations show a physical trading environment involving hardware products, timber, steel reinforcement, bamboo and local delivery activity. The business serves construction customers through direct interaction and product availability at its operating location, but it has little or no functional online marketplace through which customers can browse verified stock and place structured orders.")
h2("2.2 Current Method of Operation")
para("Nana Attakorah II Ventures currently depends substantially on physical visits, telephone calls and informal messaging. Product information, quotations and stock confirmation can be handled separately, with limited visibility after a customer makes a request. Stock may be understood by physically checking the shop or yard, and delivery arrangements may rely on direct communication with local transport operators. This method works for nearby customers but limits reach, consistency and transaction traceability.")
h2("2.3 Challenges Identified")
compact_lines([
    "1. Potential customers cannot reliably browse Nana Attakorah II Ventures' current products online.",
    "2. Informal stock confirmation can become stale and cause failed or delayed purchases.",
    "3. Quotations, order decisions, delivery events and cash evidence are difficult to audit.",
    "4. The business needs controlled staff access without exposing another supplier's private records.",
    "5. Growth beyond the immediate locality requires a trusted, searchable digital presence.",
])
h2("2.4 Justification for an Online Marketplace")
para("Onboarding Nana Attakorah II Ventures onto BuildMate gives the business a persistent digital presence without requiring a separate platform that it must operate alone. BuildMate links the public marketplace to a protected tenant workspace for the business's products, staff, branches, inventory and orders. Eligible offers appear only after supplier approval and valid stock checks. The multi-tenant model also allows BuildMate to host additional suppliers using the same application while row-level policies and permission-aware functions keep each organisation's operational data separate.")

# Page 7
new_page()
h1("Chapter Two: Business Process Requirements")
h2("2.5 Required E-Commerce Capabilities")
para("The capstone brief requires product or service browsing, search, detail viewing, cart management and online ordering. BuildMate implements these minimum functions and extends them with quotation requests, fulfilment selection, material calculators, professional-service discovery, customer workspaces and support tickets.")
h2("2.6 Business Management Requirements")
para("Within its supplier tenant, Nana Attakorah II Ventures can manage products, prices, branches, receipts, inventory, orders, quotations, staff and settlements. Its 13 evidence-backed listings remain drafts and none is marketplace-eligible until specification, current price, stock and publication readiness are confirmed. Sensitive changes use controlled status transitions and audit trails.")
h2("2.7 Inventory Integrity")
para("BuildMate distinguishes exact quantity, status-only and confirmation-required inventory. Exact availability is on-hand less reserved quantity; purchasable offers require an approved active supplier, published active listing, positive price, branch assignment and legitimate stock. The ledger records receipts and stock changes as immutable movements, with weighted-average cost, generated GRN or opening references and idempotency keys.")
h2("2.8 Catalogue Governance")
para("The canonical catalogue contains 15 customer-facing categories, 36 subcategories, 129 active products and 15 variants. Search aliases support Ghanaian trade terms and BOQ descriptions. Category media is database metadata: each top-level category has its own image path and alternative text, a unique index prevents reuse, and missing media displays a neutral placeholder.")
h2("2.9 Order and Fulfilment Control")
para("Orders use explicit acknowledgement, acceptance, preparation and dispatch states. Delivery records assignment, pickup, transit and customer confirmation; cash evidence follows fulfilment.")
h2("2.10 Support and Platform Operations")
para("A unified support centre gives customers, suppliers, drivers and service providers a controlled escalation route. Generated ticket numbers, permissions, assignments, public messages and internal notes support resolution and audit.")
h2("2.11 Conceptual Framework")
callout("Input-process-output", "Inputs include product data, stock, customer requirements and staff decisions. Processes include verification, search, quotation, checkout, ledger posting, fulfilment and support. Outputs include eligible offers, auditable orders, stock balances, notifications and operational reports.")

# Page 8
new_page()
h1("Chapter Three: System Design")
h2("3.1 Use Case Design")
para("The primary actors are customers, Nana Attakorah II Ventures acting as a supplier tenant, drivers and platform administrators. Other suppliers can be onboarded as separate tenants, and service providers and customer organisations use specialised variants of the authenticated access model. The use cases below show how public discovery connects to protected commercial and operational actions.")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.keep_with_next = True
add_picture(p.add_run(), ROOT / "use_case_diagram.png", Inches(6.7), "BuildMate Ghana use case diagram showing customer, supplier, driver and administrator interactions")
c = doc.add_paragraph("Figure 3.1: BuildMate Ghana use case diagram", style="Caption")
c.alignment = WD_ALIGN_PARAGRAPH.CENTER
h2("3.2 User Interface Design")
para("The interface uses a responsive public shell and role-specific dashboards. Shared navigation, page headers, status badges, metric cards, forms, confirmation dialogs and mobile breakpoints provide consistency while permissions determine which actions appear.")

# Page 9
new_page()
h1("Chapter Three: Database and Architecture Design")
h2("3.3 Database Design")
para("Core entities include profiles, organisations, memberships, roles, permissions, hierarchical categories, products, variants, supplier listings, branches, warehouses, inventory balances, movements, receipts, quote requests, orders, order items, deliveries, cash payments, service providers and support tickets. Category records are the source of truth for names, ordering, descriptions, image paths and accessible alternative text. UUID keys connect operational records, while generated human-readable numbers support business communication.")
h2("3.4 System Architecture")
para("The presentation layer uses Next.js 16, React 19 and Tailwind CSS. Server functions and route handlers perform application validation and orchestration. Supabase authentication and PostgreSQL provide identity, relational storage, row-level security, transactions, triggers and permission-aware RPCs. Vercel hosts the web application while the linked Supabase project provides managed data services.")
h2("3.5 Security Design")
para("Nana Attakorah II Ventures is represented by its own supplier organisation record. Its staff access the tenant through active organisation memberships, role permissions and branch or warehouse assignments. The same model can host other suppliers, but row-level policies and permission helpers restrict each membership to authorised organisation data. Direct table writes are revoked for transactional workflows, platform administration is independently permissioned, and public policies expose only eligible offers from approved active suppliers.")
h2("3.6 Transaction and Audit Design")
para("Inventory and commerce operations use row locks, unique constraints, idempotency keys and append-only events. Receipt references are generated from database sequences, preventing JavaScript-only numbering races. Audit records identify actors, entities, actions and before/after context. Corrections create compensating movements instead of rewriting historical receipts.")
h2("3.7 Verification Strategy")
compact_lines([
    "Type safety: TypeScript compiler with no emitted output.",
    "Code quality: ESLint with zero permitted warnings.",
    "Regression: Vitest security, RLS, inventory, ordering and UI assertions.",
    "Browser evidence: desktop and mobile Playwright checks plus live screenshots.",
    "Deployment readiness: optimized Next.js production build.",
])
callout("Latest verified result", "All 312 automated tests passed; TypeScript passed; ESLint passed with zero warnings; the production build completed successfully. The deployed catalogue exposes 15 active top-level categories with 15 distinct database-assigned image paths.")

# Page 10
new_page()
h1("Chapter Four: System Implementation")
h2("4.0 Technologies Used")
para("BuildMate was implemented with Next.js 16, React 19, TypeScript, Tailwind CSS, Supabase and PostgreSQL. Vitest, Playwright, ESLint and the TypeScript compiler support quality assurance, while Vercel provides web deployment.")
h2("4.1 Field Study: Stakeholder Engagement")
para("The project team visited Nana Attakorah II Ventures and its local building-materials trading environment to observe how merchants, customers and field operators discuss stock, product choice and delivery. The photographs document the real business context that informed the supplier-tenant requirements; no individual identities are inferred beyond what is visibly shown.")
field_photo_grid([
    ("WhatsApp Image 2026-08-08 at 7.37.42 PM.jpeg", "Figure 4.1: Stakeholder discussion at a retail shop", "Group discussion inside a local hardware retail shop"),
    ("WhatsApp Image 2026-08-08 at 7.37.44 PM.jpeg", "Figure 4.2: Follow-up engagement with shop participants", "Second stakeholder discussion inside the same local hardware shop"),
    ("WhatsApp Image 2026-08-08 at 6.54.18 PM.jpeg", "Figure 4.3: Field team at a timber market", "Three field-study participants standing near stacked timber"),
    ("WhatsApp Image 2026-08-08 at 6.54.21 PM.jpeg", "Figure 4.4: Supplier conversation among timber stock", "Field discussion beside stacked timber inventory"),
])

# Page 11
new_page()
h1("Chapter Four: Field Study of Supplier Operations")
h2("4.2 Premises, Stock and Local Transport")
para("The selected business operates within an environment of visible yard stock, small retail spaces and local delivery vehicles. This reinforces the need for Nana Attakorah II Ventures to maintain branch-aware stock records, product categorisation and clear fulfilment choices inside its BuildMate tenant.")
field_photo_grid([
    ("IMG_20260808_160546_418.jpg (1).jpeg", "Figure 4.5: Reinforcement steel stock", "Bundles of reinforcement steel and coiled steel at a supplier yard"),
    ("WhatsApp Image 2026-08-08 at 6.54.13 PM (1).jpeg", "Figure 4.6: Bent reinforcement bars", "Stacks of bent reinforcement bars stored outdoors"),
    ("WhatsApp Image 2026-08-08 at 6.54.28 PM.jpeg", "Figure 4.7: Local timber delivery vehicle", "Small delivery truck loaded with timber at a supplier yard"),
    ("WhatsApp Image 2026-08-08 at 6.54.30 PM.jpeg", "Figure 4.8: Finished timber products", "Supplier representative standing beside stacked wooden doors and boards"),
    ("WhatsApp Image 2026-08-08 at 6.54.32 PM.jpeg", "Figure 4.9: Roadside hardware storefront", "Local hardware storefront displaying paint and construction supplies"),
])

# Page 12
new_page()
h1("Chapter Four: Field Study of Product Variety")
h2("4.3 Timber, Bamboo and Hardware Merchandising")
para("The field evidence shows varied stock units and storage conditions: timber sold in lengths and boards, bamboo poles stored outdoors, and smaller hardware items packed densely inside retail shops. The system therefore needs flexible product units, searchable categories, photographs, stock status and supplier-specific offers.")
field_photo_grid([
    ("WhatsApp Image 2026-08-08 at 6.54.19 PM (1).jpeg", "Figure 4.10: Small hardware retailer", "Shopkeeper standing inside a densely stocked hardware shop"),
    ("WhatsApp Image 2026-08-08 at 6.54.42 PM.jpeg", "Figure 4.11: Sawn timber inventory", "Sawn timber arranged by length and section under a covered shed"),
    ("WhatsApp Image 2026-08-08 at 6.54.48 PM.jpeg", "Figure 4.12: Bamboo pole stock", "Large outdoor stacks of bamboo poles"),
    ("WhatsApp Image 2026-08-08 at 6.54.57 PM.jpeg", "Figure 4.13: Assorted hardware products", "Interior display of adhesives, tools, fasteners and roofing materials"),
    ("WhatsApp Image 2026-08-08 at 6.54.58 PM.jpeg", "Figure 4.14: Packaged fasteners and sealants", "Boxes of nails, door-frame products and silicone sealant in a retail shop"),
])

# Page 13
new_page()
h1("Chapter Four: Public Marketplace Implementation")
h2("4.4 Entry Point and Product Discovery")
system_snapshot("01-homepage.png", "Figure 4.15: BuildMate Ghana Homepage", "The homepage provides search, location capture, shopping and assisted-procurement entry points. Its category section is now populated from active top-level category metadata.")
system_snapshot("02-marketplace.png", "Figure 4.16: Building Materials Marketplace", "The offer marketplace applies supplier approval, publication, branch, price and stock eligibility. A zero-offer state does not suppress the separate 15-category catalogue.")

# Page 14
new_page()
h1("Chapter Four: Services and Planning Tools")
h2("4.5 Professional Discovery and Estimation")
system_snapshot("03-services.png", "Figure 4.17: Construction Services Marketplace", "Customers can filter verified construction professionals by category, region and availability.")
system_snapshot("04-calculators.png", "Figure 4.18: Material Calculators", "Preliminary estimators support common construction quantities while preserving a clear professional-validation disclaimer.")

# Page 15
new_page()
h1("Chapter Four: Quotation and Checkout Workflows")
h2("4.6 Structured Procurement and Fulfilment")
system_snapshot("05-request-quote.png", "Figure 4.19: Competitive Quotation Request Interface", "The RFQ form captures project location, delivery date, material requirements and customer instructions.")
system_snapshot("06-cart-checkout.png", "Figure 4.20: Shopping Cart and Fulfilment Interface", "The cart separates delivery from supplier pickup and communicates cash-on-delivery or cash-on-pickup terms.")

# Page 16
new_page()
h1("Chapter Four: Responsive and Administrative Interfaces")
h2("4.7 Responsive Access and Platform Operations")
pair = doc.add_table(rows=1, cols=2)
set_table_geometry(pair, [3300, 6060])
pair.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell in pair.rows[0].cells:
    cell._tc.get_or_add_tcPr().append(OxmlElement("w:tcBorders"))
left, right = pair.rows[0].cells
left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_picture(left.paragraphs[0].add_run(), SHOTS / "07-mobile-homepage.png", Inches(1.72), "Responsive mobile view of the BuildMate Ghana homepage")
lp = left.add_paragraph("Figure 4.21: Responsive Mobile Homepage", style="Caption")
lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
left.add_paragraph("The 390-pixel capture shows legible mobile navigation, procurement actions and support without horizontal overflow.", style="Figure Explanation")
right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_picture(right.paragraphs[0].add_run(), SHOTS / "08-admin-dashboard.png", Inches(4.05), "BuildMate Ghana platform administration dashboard")
rp = right.add_paragraph("Figure 4.22: Platform Administration Dashboard", style="Caption")
rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
right.add_paragraph("The live dashboard displays customer, supplier, order, revenue, support and risk indicators, demonstrating operational governance without exposing credentials or developer tooling.", style="Figure Explanation")
para("Together, Figures 4.21 and 4.22 demonstrate the system's breadth: responsive customer-facing access at one end and permission-controlled platform supervision at the other.")
h2("4.8 Description of Implemented Functionalities")
para("Customer functionality includes product browsing, search and filtering, product-detail and supplier-offer comparison, shopping-cart management, delivery or supplier-pickup selection, online order placement, competitive quotation requests, service-provider discovery, material calculators, order tracking and support tickets.")
para("Business and platform functionality includes supplier product and price management, branch-aware inventory, opening stock and receipts, movement history, staff permissions, quotations, fulfilment, delivery, settlements, verification, reporting, audit logs and support. Migrations 202608220071-202608220074 also delivered the complete database-driven category catalogue, current versioned category imagery and Nana's draft-only evidence safeguards without fabricating stock, cost, warehouse or publication state.")

# Page 17
new_page()
h1("Chapter Five: Conclusion and Recommendations")
h2("5.1 Summary of Findings")
para("Repository, field-study and browser evidence confirm that BuildMate can support Nana Attakorah II Ventures as a supplier tenant: the complete 15-category public catalogue and quotation interfaces are connected to protected supplier tools, branch-aware inventory, explicit order and delivery states, support operations, audit trails and platform administration. Organisation-level policies separate the selected business from other tenants, while marketplace eligibility rules prevent its 13 inactive draft listings, unavailable stock or unapproved offers from being advertised.")
h2("5.2 Challenges Encountered")
para("The main implementation challenges were coordinating multi-role navigation, enforcing branch-aware supplier access, preserving inventory integrity under concurrent operations, aligning order and delivery state machines, and separating a complete canonical catalogue from marketplace offers that must satisfy stricter eligibility rules. Category media also had to be moved from repeated front-end fallbacks to database-owned, unique and accessible metadata. During the original evidence capture, production data exposed no eligible listings and the authenticated supplier/customer session expired, limiting the available protected screenshots.")
h2("5.3 Quality Results")
para("TypeScript, ESLint and production compilation completed successfully, and all 312 automated tests passed. Migration verification confirmed 15 active top-level categories, 36 active subcategories, 129 active catalogue products, 15 product variants and 13 Nana Attakorah II Ventures draft listings. All 15 top-level category image paths are assigned and distinct, while zero Nana listings are accidentally marketplace-eligible. Production HTML also exposed the four latest versioned category-image paths after deployment.")
h2("5.4 Conclusion")
para("The project demonstrates a practical route for digitising Nana Attakorah II Ventures without building an isolated single-business website. By onboarding the business as a supplier tenant, BuildMate combines public marketplace reach with private organisation controls: stock is ledger-based, orders transition explicitly, staff permissions are enforced below the interface and operational changes are audited. The approach can support the selected business now and additional suppliers later without mixing their confidential records.")
h2("5.5 Recommendations for Future Improvement")
compact_lines([
    "1. Continue applying future forward-only migrations through the controlled deployment pipeline used for migrations 202608220070-202608220074.",
    "2. Obtain Nana Attakorah II Ventures' confirmation for product specifications, current prices, stock and opening balances.",
    "3. Publish representative offers only after supplier approval, branch assignment, positive price and eligible stock are verified.",
    "4. Add isolated end-to-end fixtures for marketplace, checkout and inventory acceptance tests.",
    "5. Continue accessibility, performance, backup and incident-response exercises before broad launch.",
])

# Page 18
new_page()
h1("References and Project-Brief Alignment")
h2("References")
compact_lines([
    "Cyber Security Authority Ghana. (2012). Data Protection Act, 2012 (Act 843). https://csa.gov.gh/resources/Data_Protection_Act_2012.pdf",
    "Next.js. (2026). App Router documentation. https://nextjs.org/docs/app",
    "Supabase. (2026). Row Level Security. https://supabase.com/docs/guides/database/postgres/row-level-security",
    "Supabase. (2026). Securing the Data API. https://supabase.com/docs/guides/api/securing-your-api",
    "BuildMate Ghana repository. (2026). Application routes, tests, category assets and Supabase migrations inspected through commit 9b19f8f.",
    "BuildMate Ghana deployment. (2026). Live system evidence reviewed at https://buildmate-six.vercel.app/ through 22 August 2026.",
])
h2("Alignment with the Supplied Capstone Brief")
para("The supplied E-Commerce capstone brief requires a real local business with little or no online marketplace and a web-based solution through which customers can browse products and place orders. Nana Attakorah II Ventures satisfies the business-selection requirement, while BuildMate satisfies and extends the minimum functional requirements.")
compact_lines([
    "Customer module: view and search products, inspect product details, add items to a cart and place orders online.",
    "Business management: add products, edit product information, control publication or archival, and view customer orders.",
    "Required deliverables: a professional project report, a functional marketplace, and a presentation with a live demonstration.",
    "Required report structure: Introduction, Business Analysis, System Design, System Implementation, and Conclusion and Recommendations.",
    "Submission deadline stated in the brief: Saturday, 22 August 2026.",
])
h2("Assessment Criteria")
assessment = doc.add_table(rows=1, cols=2)
assessment.style = "Table Grid"
assessment.rows[0].cells[0].text = "Component"
assessment.rows[0].cells[1].text = "Marks"
for cell in assessment.rows[0].cells:
    set_cell_shading(cell, LIGHT)
    cell.paragraphs[0].runs[0].bold = True
for component, marks in [
    ("Business Selection and Analysis", "5"),
    ("Application of E-Commerce Concepts", "10"),
    ("System Design", "5"),
    ("System Functionality", "10"),
    ("Documentation (Project Report)", "20"),
    ("Presentation and Demonstration", "10"),
    ("Total", "60"),
]:
    cells = assessment.add_row().cells
    cells[0].text = component
    cells[1].text = marks
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
set_table_geometry(assessment, [7800, 1560])
callout("Evidence statement", "Evidence comprises the Regent brief, 14 field photographs, eight system screenshots, repository inspection, migrations 202608220070-202608220074 and recorded quality checks. No screenshots, stock, prices, credentials or personal details were fabricated.")

# Apply header and metadata
header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = header.add_run("BUILDMATE GHANA  |  CAPSTONE REPORT")
run.font.name = "Times New Roman"
run.font.size = Pt(8)
run.font.bold = True
run.font.color.rgb = GREEN

props = doc.core_properties
props.title = "BuildMate Ghana Capstone Project Report"
props.subject = "Digitising Nana Attakorah II Ventures through a secure multi-tenant construction marketplace"
props.author = "David Selorm Gabion; Sualiha Nasomah Zakari; Charles Nanor; Shadrack Ofori; Mary Sarpong"
props.keywords = "BuildMate, Nana Attakorah II Ventures, multi-tenant marketplace, Ghana, construction procurement, inventory, Supabase, Next.js"

doc.save(OUT)
print(OUT)
